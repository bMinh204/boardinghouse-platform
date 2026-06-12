from pathlib import Path

from docx import Document


SOURCE = Path(r"E:\DATN\project\boardinghouse_platform\docs\hop-dong-thue-nha-tro-template.docx")
OUTPUT = Path(r"E:\DATN\project\boardinghouse_platform\src\main\resources\templates\hop-dong-thue-nha-tro.docx")


def replace_paragraph(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


document = Document(SOURCE)
replacements = {
    5: "Hôm nay, {{CONTRACT_DATE_TEXT}}, tại {{CONTRACT_LOCATION}}. Chúng tôi ký tên dưới đây gồm có:",
    8: "Ông/bà: {{LANDLORD_NAME}}",
    9: "CMND/CCCD số: {{LANDLORD_CCCD}}    Ngày sinh: {{LANDLORD_DOB}}",
    10: "Thường trú tại: {{LANDLORD_ADDRESS}}    Số điện thoại: {{LANDLORD_PHONE}}",
    12: "Ông/bà: {{TENANT_NAME}}",
    13: "CMND/CCCD số: {{TENANT_CCCD}}    Ngày sinh: {{TENANT_DOB}}",
    14: "Thường trú tại: {{TENANT_ADDRESS}}    Số điện thoại: {{TENANT_PHONE}}",
    17: (
        "Bên A cho Bên B thuê 01 phòng trọ {{ROOM_TITLE}} tại địa chỉ {{ROOM_ADDRESS}}. "
        "Diện tích {{ROOM_SIZE}} m², sức chứa tối đa {{ROOM_CAPACITY}} người. "
        "Thời hạn thuê là {{DURATION_MONTHS}} tháng, từ {{START_DATE}} đến {{END_DATE}}. "
        "Giá thuê: {{RENT}} đồng/tháng (Bằng chữ: {{RENT_WORDS}}). "
        "Chưa bao gồm chi phí điện, nước và các dịch vụ khác."
    ),
    22: (
        "Bên B đặt cọc số tiền {{DEPOSIT}} đồng (Bằng chữ: {{DEPOSIT_WORDS}}). "
        "Tiền thuê được thanh toán theo kỳ: {{PAYMENT_CYCLE}}."
    ),
    24: (
        "Chỉ sử dụng phòng trọ vào mục đích ở, với số lượng tối đa không quá "
        "{{ROOM_CAPACITY}} người; không chứa thiết bị gây cháy nổ, hàng cấm; "
        "cung cấp giấy tờ tùy thân để đăng ký tạm trú và tuân thủ quy định pháp luật."
    ),
    28: (
        "Sau thời hạn thuê {{DURATION_MONTHS}} tháng, nếu Bên B có nhu cầu tiếp tục thuê, "
        "hai bên sẽ thương lượng để gia hạn hợp đồng."
    ),
    30: "                Bên B                                                        Bên A",
    31: (
        "        (Ký, ghi rõ họ tên)                                      (Ký, ghi rõ họ tên)\n"
        "        {{TENANT_NAME}}                                           {{LANDLORD_NAME}}"
    ),
}

for index, text in replacements.items():
    replace_paragraph(document.paragraphs[index], text)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
document.save(OUTPUT)
print(OUTPUT)
